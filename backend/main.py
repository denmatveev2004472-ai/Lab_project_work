from fastapi import FastAPI, HTTPException, Query, Depends, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List
from sqlalchemy.orm import Session
from sqlalchemy import text
from sqlalchemy.exc import IntegrityError
from io import BytesIO
from openpyxl import Workbook
from docx import Document
from datetime import datetime, timedelta, date
import os
import json
from dotenv import load_dotenv

from database import get_db, init_db
from models import Location, Item

load_dotenv()

app = FastAPI(title="Lab Inventory API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

ALLOWED_TYPES = {"reagent", "equipment", "consumable", "furniture", "other"}

# ── Pydantic-схемы ────────────────────────────────────────────────────────────

class LocationIn(BaseModel):
    room: str
    cabinet: Optional[str] = None
    shelf: Optional[str] = None
    slot: Optional[str] = None

class ItemIn(BaseModel):
    item_type: str = "other"
    code: Optional[str] = None
    name: Optional[str] = None
    name_ru: Optional[str] = None
    name_en: Optional[str] = None
    formula: Optional[str] = None
    cas: Optional[str] = None
    manufacturer: Optional[str] = None
    catalog_number: Optional[str] = None
    inventory_number: Optional[str] = None
    serial_number: Optional[str] = None
    registry_number: Optional[str] = None
    quantity: Optional[str] = None
    unit: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None
    source_file: Optional[str] = None
    source_sheet: Optional[str] = None
    location_id: Optional[int] = None
    document_url: Optional[str] = None
    gdrive_link: Optional[str] = None

class ItemUpdate(BaseModel):
    is_out_of_stock: Optional[bool] = None
    item_type: Optional[str] = None
    code: Optional[str] = None
    name: Optional[str] = None
    name_ru: Optional[str] = None
    name_en: Optional[str] = None
    formula: Optional[str] = None
    cas: Optional[str] = None
    manufacturer: Optional[str] = None
    catalog_number: Optional[str] = None
    inventory_number: Optional[str] = None
    serial_number: Optional[str] = None
    registry_number: Optional[str] = None
    quantity: Optional[str] = None
    unit: Optional[str] = None
    status: Optional[str] = None
    notes: Optional[str] = None
    source_file: Optional[str] = None
    source_sheet: Optional[str] = None
    location_id: Optional[int] = None
    document_url: Optional[str] = None
    gdrive_link: Optional[str] = None

class ProtocolIn(BaseModel):
    code: Optional[str] = None
    name: str
    category: str = "other"
    steps: Optional[str] = None
    notes: Optional[str] = None

class ProtocolUpdate(BaseModel):
    code: Optional[str] = None
    name: Optional[str] = None
    category: Optional[str] = None
    steps: Optional[str] = None
    notes: Optional[str] = None

class EquipmentUsageIn(BaseModel):
    user_name: str
    start_time: datetime
    end_time: datetime
    notes: Optional[str] = None
    force: bool = False  # если True — разрешаем перекрытие (конфликт)

class MaintenanceIn(BaseModel):
    date_from: date
    date_to: date
    reason: Optional[str] = None
    created_by: Optional[str] = None

# Клеточная культура — бронирование боксов
class CellUsageIn(BaseModel):
    user_name: str
    box_number: int  # 1 или 2
    start_time: datetime
    end_time: datetime
    experiment: Optional[str] = None
    force: bool = False

class CellUsageUpdate(BaseModel):
    user_name: Optional[str] = None
    box_number: Optional[int] = None
    start_time: Optional[datetime] = None
    end_time: Optional[datetime] = None
    experiment: Optional[str] = None


class InstrumentUsageIn(BaseModel):
    instrument_id: str
    username: str
    start_time: datetime
    end_time: datetime
    notes: Optional[str] = None
    force: bool = False

class InstrumentUsageUpdateIn(BaseModel):
    username: Optional[str] = None
    start_time: Optional[datetime] = None
    end_time: Optional[datetime] = None
    notes: Optional[str] = None

class InstrumentMaintenanceIn(BaseModel):
    instrument_id: str
    date_from: date
    date_to: date
    reason: Optional[str] = None
    created_by: Optional[str] = None

# ── Вспомогательные функции ───────────────────────────────────────────────────

def normalize_item_type(value: Optional[str]) -> str:
    val = (value or "other").strip().lower()
    return val if val in ALLOWED_TYPES else "other"

def update_search_vector(db: Session, item_id: int):
    db.execute(
        text("""
UPDATE items
SET
  updated_at = now(),
  search_vector =
    to_tsvector('russian',
      coalesce(name,'')    || ' ' ||
      coalesce(name_ru,'') || ' ' ||
      coalesce(notes,'')) ||
    to_tsvector('english',
      coalesce(name,'')             || ' ' ||
      coalesce(name_en,'')           || ' ' ||
      coalesce(formula,'')           || ' ' ||
      coalesce(cas,'')               || ' ' ||
      coalesce(internal_code,'')     || ' ' ||
      coalesce(manufacturer,'')      || ' ' ||
      coalesce(catalog_number,'')    || ' ' ||
      coalesce(inventory_number,'')  || ' ' ||
      coalesce(serial_number,'')     || ' ' ||
      coalesce(registry_number,'')   || ' ' ||
      coalesce(status,'')            || ' ' ||
      coalesce(source_file,'')       || ' ' ||
      coalesce(source_sheet,''))
WHERE id = :item_id
        """),
        {"item_id": item_id},
    )

def get_or_create_location(db: Session, room: Optional[str], cabinet: Optional[str],
                            shelf: Optional[str], slot: Optional[str]) -> Optional[int]:
    if not room:
        return None
    row = db.execute(
        text("""
INSERT INTO locations (room, cabinet, shelf, slot)
VALUES (:room, :cabinet, :shelf, :slot)
ON CONFLICT (room, cabinet, shelf, slot)
DO UPDATE SET room = EXCLUDED.room
RETURNING id
        """),
        {"room": room, "cabinet": cabinet or "-", "shelf": shelf or "-", "slot": slot or "-"},
    ).fetchone()
    return row[0] if row else None

def _export_excel(rows, headers, filename):
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for r in rows:
        ws.append(list(r))
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

def _export_word_table(rows, headers, title, filename):
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = str(val) if val is not None else ""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )

def log_activity(db: Session, action: str, entity_type: str, entity_id, name: str, details: str = ""):
    db.execute(
        text("""
INSERT INTO activity_log (action, entity_type, entity_id, name, details)
VALUES (:action, :entity_type, :entity_id, :name, :details)
        """),
        {
            "action": action,
            "entity_type": entity_type,
            "entity_id": str(entity_id) if entity_id is not None else None,
            "name": name or "—",
            "details": details or "",
        },
    )

def period_range(period: str, offset: int = 0):
    now = datetime.now()
    period = (period or "week").lower()
    if period == "day":
        base = now + timedelta(days=offset)
        start = base.replace(hour=0, minute=0, second=0, microsecond=0)
        end = start + timedelta(days=1)
    elif period == "month":
        month_index = now.month - 1 + offset
        year = now.year + month_index // 12
        month = month_index % 12 + 1
        start = datetime(year, month, 1)
        end = datetime(year + (1 if month == 12 else 0), 1 if month == 12 else month + 1, 1)
    elif period == "year":
        year = now.year + offset
        start = datetime(year, 1, 1)
        end = datetime(year + 1, 1, 1)
    else:  # week
        base = now + timedelta(weeks=offset)
        start = (base - timedelta(days=base.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
        end = start + timedelta(days=7)
    return start, end

def ensure_extra_tables(db: Session):
    db.execute(text("""
CREATE TABLE IF NOT EXISTS activity_log (
    id SERIAL PRIMARY KEY,
    action TEXT NOT NULL,
    entity_type TEXT NOT NULL,
    entity_id TEXT,
    name TEXT,
    details TEXT,
    created_at TIMESTAMP NOT NULL DEFAULT now()
)
    """))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_activity_log_created_at ON activity_log (created_at)"))

    db.execute(text("""
CREATE TABLE IF NOT EXISTS equipment_usage (
    id SERIAL PRIMARY KEY,
    item_id INTEGER NOT NULL REFERENCES items(id) ON DELETE CASCADE,
    user_name TEXT NOT NULL,
    start_time TIMESTAMP NOT NULL,
    end_time TIMESTAMP NOT NULL,
    notes TEXT,
    created_at TIMESTAMP NOT NULL DEFAULT now()
)
    """))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_equipment_usage_item ON equipment_usage (item_id, start_time)"))

    db.execute(text("CREATE EXTENSION IF NOT EXISTS btree_gist"))
    db.execute(text("""
DO $$
BEGIN
    IF NOT EXISTS (
        SELECT 1 FROM pg_constraint WHERE conname = 'equipment_usage_no_overlap'
    ) THEN
        ALTER TABLE equipment_usage
            ADD CONSTRAINT equipment_usage_no_overlap
            EXCLUDE USING gist (
                item_id WITH =,
                tsrange(start_time, end_time) WITH &&
            );
    END IF;
END$$;
    """))

    db.execute(text("""
CREATE TABLE IF NOT EXISTS maintenance (
    id SERIAL PRIMARY KEY,
    item_id INTEGER NOT NULL REFERENCES items(id) ON DELETE CASCADE,
    date_from DATE NOT NULL,
    date_to DATE NOT NULL,
    reason TEXT,
    created_by TEXT,
    created_at TIMESTAMP NOT NULL DEFAULT now()
)
    """))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_maintenance_item ON maintenance (item_id)"))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_maintenance_dates ON maintenance (date_from, date_to)"))

    # Клеточная культура — отдельная таблица бронирования боксов
    db.execute(text("""CREATE TABLE IF NOT EXISTS cell_usage (
        id SERIAL PRIMARY KEY,
        box_number INTEGER NOT NULL CHECK (box_number IN (1, 2)),
        username TEXT NOT NULL,
        start_time TIMESTAMP NOT NULL,
        end_time TIMESTAMP NOT NULL,
        experiment TEXT,
        created_at TIMESTAMP NOT NULL DEFAULT now()
    )"""))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_cell_usage_box ON cell_usage (box_number, start_time)"))

    db.execute(text("""CREATE TABLE IF NOT EXISTS instrument_usage (
        id SERIAL PRIMARY KEY,
        instrument_id TEXT NOT NULL,
        username TEXT NOT NULL,
        start_time TIMESTAMP NOT NULL,
        end_time TIMESTAMP NOT NULL,
        notes TEXT,
        created_at TIMESTAMP NOT NULL DEFAULT now()
    )"""))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_instrument_usage_instr ON instrument_usage (instrument_id, start_time)"))

    db.execute(text("""CREATE TABLE IF NOT EXISTS instrument_maintenance (
        id SERIAL PRIMARY KEY,
        instrument_id TEXT NOT NULL,
        date_from DATE NOT NULL,
        date_to DATE NOT NULL,
        reason TEXT,
        created_by TEXT,
        created_at TIMESTAMP NOT NULL DEFAULT now()
    )"""))
    db.execute(text("CREATE INDEX IF NOT EXISTS idx_instrument_maintenance_instr ON instrument_maintenance (instrument_id)"))
    db.commit()

# ── Жизненный цикл ────────────────────────────────────────────────────────────

@app.on_event("startup")
def on_startup():
    init_db()
    db = next(get_db())
    try:
        ensure_extra_tables(db)
    finally:
        db.close()

# ── Эндпоинты: health / search / stats / rooms ────────────────────────────────

@app.get("/api/health")
def health():
    return {"ok": True}

@app.get("/api/search")
def search(
    q: str = Query(default=""),
    room: str = Query(default=""),
    cabinet: str = Query(default=""),
    item_type: str = Query(default=""),
    source_file: str = Query(default=""),
    limit: int = Query(default=300, ge=1, le=1000),
    db: Session = Depends(get_db),
):
    params = {
        "q": q,
        "like": f"%{q}%",
        "room": room,
        "cabinet": cabinet,
        "item_type": normalize_item_type(item_type) if item_type else "",
        "source_file": source_file,
        "limit": limit,
    }
    rows = db.execute(
        text("""
SELECT
  i.id, i.item_type, i.internal_code AS code,
  i.name, i.name_ru, i.name_en,
  i.formula, i.cas, i.manufacturer,
  i.catalog_number, i.inventory_number, i.serial_number,
  i.registry_number, i.quantity, i.unit,
  i.status, i.notes, i.source_file, i.source_sheet,
  i.location_id, i.document_url, i.gdrive_link,
  l.room, l.cabinet, l.shelf, l.slot,
  coalesce(m.in_maintenance, false) AS in_maintenance
FROM items i
LEFT JOIN locations l ON i.location_id = l.id
LEFT JOIN (
  SELECT item_id,
         bool_or(CURRENT_DATE BETWEEN date_from AND date_to) AS in_maintenance
  FROM maintenance
  GROUP BY item_id
) m ON m.item_id = i.id
WHERE (:q = ''
    OR i.search_vector @@ plainto_tsquery('russian', :q)
    OR i.search_vector @@ plainto_tsquery('english', :q)
    OR coalesce(i.name,'')           ILIKE :like
    OR coalesce(i.name_ru,'')        ILIKE :like
    OR coalesce(i.name_en,'')        ILIKE :like
    OR coalesce(i.formula,'')        ILIKE :like
    OR coalesce(i.cas,'')            ILIKE :like
    OR coalesce(i.internal_code,'')  ILIKE :like
    OR coalesce(i.inventory_number,'') ILIKE :like
    OR coalesce(i.catalog_number,'') ILIKE :like
    OR coalesce(i.serial_number,'')  ILIKE :like
    OR coalesce(i.registry_number,'') ILIKE :like
    OR coalesce(i.manufacturer,'')   ILIKE :like
    OR coalesce(i.notes,'')          ILIKE :like
)
AND (:room = '' OR coalesce(l.room,'') = :room)
AND (:cabinet = '' OR coalesce(l.cabinet,'') = :cabinet)
AND (:item_type = '' OR coalesce(i.item_type,'') = :item_type)
AND (:source_file = '' OR coalesce(i.source_file,'') = :source_file)
ORDER BY
  coalesce(l.room,'ZZZ'),
  coalesce(l.cabinet,'ZZZ'),
  coalesce(i.item_type,'zzz'),
  coalesce(i.name, i.name_ru)
LIMIT :limit
        """),
        params,
    ).fetchall()
    results = [dict(r._mapping) for r in rows]
    return {"count": len(results), "results": results}

@app.get("/api/stats")
def stats(db: Session = Depends(get_db)):
    total = db.execute(text("SELECT COUNT(*) FROM items")).scalar()
    by_type = db.execute(text("SELECT item_type, COUNT(*) AS n FROM items GROUP BY item_type ORDER BY item_type")).fetchall()
    by_room = db.execute(text("""
SELECT coalesce(l.room,'') AS room, COUNT(*) AS n
FROM items i
LEFT JOIN locations l ON i.location_id = l.id
GROUP BY coalesce(l.room,'')
ORDER BY room
    """)).fetchall()
    files = db.execute(text("SELECT source_file, COUNT(*) AS n FROM items GROUP BY source_file ORDER BY source_file")).fetchall()
    protocols_total = db.execute(text("SELECT COUNT(*) FROM protocols")).scalar()
    return {
        "total": total,
        "protocols_total": protocols_total,
        "by_type": [dict(r._mapping) for r in by_type],
        "by_room": [dict(r._mapping) for r in by_room],
        "by_source_file": [dict(r._mapping) for r in files],
    }

@app.get("/api/rooms")
def rooms(item_type: str = Query(default=""), db: Session = Depends(get_db)):
    params = {"item_type": item_type}
    where = "AND i.item_type = :item_type" if item_type else ""
    rows = db.execute(text(f"""
SELECT l.room, COUNT(i.id) AS items_count
FROM locations l
LEFT JOIN items i ON i.location_id = l.id
WHERE 1=1 {where}
GROUP BY l.room
HAVING COUNT(i.id) > 0
ORDER BY l.room
    """), params).fetchall()
    return [dict(r._mapping) for r in rows]

@app.get("/api/locations/summary")
def locations_summary(room: str = Query(default=""), db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT l.cabinet, COUNT(i.id) AS items_count
FROM locations l
LEFT JOIN items i ON i.location_id = l.id
WHERE :room = '' OR l.room = :room
GROUP BY l.cabinet
ORDER BY l.cabinet
    """), {"room": room}).fetchall()
    return [dict(r._mapping) for r in rows]

@app.get("/api/locations")
def get_locations(room: str = Query(default=""), db: Session = Depends(get_db)):
    if room:
        rows = db.execute(text("SELECT * FROM locations WHERE room = :room ORDER BY room, cabinet, shelf, slot"), {"room": room}).fetchall()
    else:
        rows = db.execute(text("SELECT * FROM locations ORDER BY room, cabinet, shelf, slot")).fetchall()
    return [dict(r._mapping) for r in rows]

@app.post("/api/locations")
def add_location(loc: LocationIn, db: Session = Depends(get_db)):
    row = db.execute(
        text("""
INSERT INTO locations (room, cabinet, shelf, slot)
VALUES (:room, :cabinet, :shelf, :slot)
ON CONFLICT (room, cabinet, shelf, slot)
DO UPDATE SET room = EXCLUDED.room
RETURNING *
        """),
        {"room": loc.room, "cabinet": loc.cabinet or "-", "shelf": loc.shelf or "-", "slot": loc.slot or "-"},
    ).fetchone()
    db.commit()
    return dict(row._mapping)

# ── Items CRUD ────────────────────────────────────────────────────────────────

@app.get("/api/item/{item_id}")
def get_item(item_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("""
SELECT i.*, i.internal_code AS code, l.room, l.cabinet, l.shelf, l.slot
FROM items i
LEFT JOIN locations l ON i.location_id = l.id
WHERE i.id = :item_id
    """), {"item_id": item_id}).fetchone()
    if not row:
        raise HTTPException(404, "Не найдено")
    return dict(row._mapping)

@app.post("/api/item")
def add_item(item: ItemIn, db: Session = Depends(get_db)):
    item_type = normalize_item_type(item.item_type)
    row = db.execute(
        text("""
INSERT INTO items (item_type, internal_code, name, name_ru, name_en, formula, cas,
    manufacturer, catalog_number, inventory_number, serial_number,
    registry_number, quantity, unit, status, notes,
    source_file, source_sheet, location_id, document_url, gdrive_link)
VALUES (:item_type, :internal_code, :name, :name_ru, :name_en, :formula, :cas,
    :manufacturer, :catalog_number, :inventory_number, :serial_number,
    :registry_number, :quantity, :unit, :status, :notes,
    :source_file, :source_sheet, :location_id, :document_url, :gdrive_link)
RETURNING id
        """),
        {
            "item_type": item_type,
            "internal_code": item.code,
            "name": item.name,
            "name_ru": item.name_ru,
            "name_en": item.name_en,
            "formula": item.formula,
            "cas": item.cas,
            "manufacturer": item.manufacturer,
            "catalog_number": item.catalog_number,
            "inventory_number": item.inventory_number,
            "serial_number": item.serial_number,
            "registry_number": item.registry_number,
            "quantity": item.quantity,
            "unit": item.unit,
            "status": item.status,
            "notes": item.notes,
            "source_file": item.source_file,
            "source_sheet": item.source_sheet,
            "location_id": item.location_id,
            "document_url": item.document_url,
            "gdrive_link": item.gdrive_link,
        },
    ).fetchone()
    item_id = row[0]
    update_search_vector(db, item_id)
    log_activity(db, "добавлен", "item", item_id, item.name or item.code or "—", item_type)
    db.commit()
    return {"ok": True, "id": item_id}

@app.patch("/api/item/{item_id}")
def update_item(item_id: int, data: ItemUpdate, db: Session = Depends(get_db)):
    updates = {k: v for k, v in data.dict().items() if v is not None}
    if not updates:
        raise HTTPException(400, "Нет данных")
    if "item_type" in updates:
        updates["item_type"] = normalize_item_type(updates["item_type"])
    if "code" in updates:
        updates["internal_code"] = updates.pop("code")
    sets = ", ".join(f"{k} = :{k}" for k in updates)
    updates["item_id"] = item_id
    cur = db.execute(text(f"UPDATE items SET {sets} WHERE id = :item_id"), updates)
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    update_search_vector(db, item_id)
    log_name = updates.get("name")
    if not log_name:
        current = db.execute(text("SELECT name, internal_code FROM items WHERE id = :id"), {"id": item_id}).fetchone()
        log_name = current[0] or current[1] if current else "—"
    log_activity(db, "изменён", "item", item_id, log_name, "")
    db.commit()
    return {"ok": True}

@app.patch("/api/item/{item_id}/toggle-stock")
def toggle_stock(item_id: int, db: Session = Depends(get_db)):
    row = db.execute(
        text("UPDATE items SET is_out_of_stock = NOT coalesce(is_out_of_stock, false) WHERE id = :item_id RETURNING is_out_of_stock"),
        {"item_id": item_id},
    ).fetchone()
    if not row:
        raise HTTPException(404, "Не найдено")
    update_search_vector(db, item_id)
    db.commit()
    return {"ok": True, "is_out_of_stock": row[0]}

@app.delete("/api/item/{item_id}")
def delete_item(item_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT name, internal_code FROM items WHERE id = :id"), {"id": item_id}).fetchone()
    cur = db.execute(text("DELETE FROM items WHERE id = :item_id"), {"item_id": item_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    name = row[0] or row[1] or "—" if row else "—"
    log_activity(db, "удалён", "item", item_id, name, "")
    db.commit()
    return {"ok": True}

# ── Protocols CRUD ────────────────────────────────────────────────────────────

@app.get("/api/protocols")
def list_protocols(q: str = Query(default=""), category: str = Query(default=""), db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT * FROM protocols
WHERE (:q = '' OR name ILIKE :like OR coalesce(code,'') ILIKE :like OR coalesce(steps,'') ILIKE :like)
AND (:category = '' OR category = :category)
ORDER BY category, name
    """), {"q": q, "like": f"%{q}%", "category": category}).fetchall()
    return [dict(r._mapping) for r in rows]

@app.get("/api/protocols/categories")
def protocol_categories(db: Session = Depends(get_db)):
    rows = db.execute(text("SELECT category, COUNT(*) AS n FROM protocols GROUP BY category ORDER BY category")).fetchall()
    return [dict(r._mapping) for r in rows]

@app.get("/api/protocols/{protocol_id}")
def get_protocol(protocol_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT * FROM protocols WHERE id = :id"), {"id": protocol_id}).fetchone()
    if not row:
        raise HTTPException(404, "Не найдено")
    return dict(row._mapping)

@app.post("/api/protocols")
def add_protocol(p: ProtocolIn, db: Session = Depends(get_db)):
    row = db.execute(
        text("INSERT INTO protocols (code, name, category, steps, notes) VALUES (:code, :name, :category, :steps, :notes) RETURNING id"),
        p.dict(),
    ).fetchone()
    log_activity(db, "добавлен", "protocol", row[0], p.name, p.category)
    db.commit()
    return {"ok": True, "id": row[0]}

@app.patch("/api/protocols/{protocol_id}")
def update_protocol(protocol_id: int, data: ProtocolUpdate, db: Session = Depends(get_db)):
    updates = {k: v for k, v in data.dict().items() if v is not None}
    if not updates:
        raise HTTPException(400, "Нет данных")
    sets = ", ".join(f"{k} = :{k}" for k in updates)
    updates["id"] = protocol_id
    cur = db.execute(text(f"UPDATE protocols SET {sets}, updated_at = now() WHERE id = :id"), updates)
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    log_name = updates.get("name")
    if not log_name:
        current = db.execute(text("SELECT name FROM protocols WHERE id = :id"), {"id": protocol_id}).fetchone()
        log_name = current[0] if current else "—"
    log_activity(db, "изменён", "protocol", protocol_id, log_name, "")
    db.commit()
    return {"ok": True}

@app.delete("/api/protocols/{protocol_id}")
def delete_protocol(protocol_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT name FROM protocols WHERE id = :id"), {"id": protocol_id}).fetchone()
    cur = db.execute(text("DELETE FROM protocols WHERE id = :id"), {"id": protocol_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    log_activity(db, "удалён", "protocol", protocol_id, row[0] if row else "—", "")
    db.commit()
    return {"ok": True}

# ── Activity log ──────────────────────────────────────────────────────────────

@app.get("/api/activity")
def list_activity(
    period: str = Query(default="week"),
    offset: int = Query(default=0),
    limit: int = Query(default=200, ge=1, le=1000),
    db: Session = Depends(get_db),
):
    start, end = period_range(period, offset)
    rows = db.execute(text("""
SELECT id, action, entity_type, entity_id, name, details, created_at
FROM activity_log
WHERE created_at >= :start AND created_at < :end
ORDER BY created_at DESC
LIMIT :limit
    """), {"start": start, "end": end, "limit": limit}).fetchall()
    return {
        "period": period,
        "offset": offset,
        "range_start": start,
        "range_end": end,
        "results": [dict(r._mapping) for r in rows],
    }

# ── Equipment usage (аналитические приборы) ───────────────────────────────────

@app.get("/api/equipment/{item_id}/usage")
def list_equipment_usage(
    item_id: int,
    period: str = Query(default="week"),
    offset: int = Query(default=0),
    db: Session = Depends(get_db),
):
    start, end = period_range(period, offset)
    rows = db.execute(text("""
SELECT id, item_id, user_name, start_time, end_time, notes, created_at
FROM equipment_usage
WHERE item_id = :item_id AND start_time >= :start AND end_time <= :end
ORDER BY start_time
    """), {"item_id": item_id, "start": start, "end": end}).fetchall()
    return {
        "period": period,
        "offset": offset,
        "range_start": start,
        "range_end": end,
        "results": [dict(r._mapping) for r in rows],
    }

@app.get("/api/equipment/usage/all")
def list_all_usage(
    period: str = Query(default="week"),
    offset: int = Query(default=0),
    db: Session = Depends(get_db),
):
    start, end = period_range(period, offset)
    rows = db.execute(text("""
SELECT u.id, u.item_id, u.user_name, u.start_time, u.end_time, u.notes, u.created_at,
       i.name as item_name
FROM equipment_usage u
JOIN items i ON u.item_id = i.id
WHERE u.start_time >= :start AND u.end_time <= :end
ORDER BY u.start_time
    """), {"start": start, "end": end}).fetchall()
    return {
        "period": period,
        "offset": offset,
        "range_start": start,
        "range_end": end,
        "results": [dict(r._mapping) for r in rows],
    }

@app.post("/api/equipment/{item_id}/usage")
def add_equipment_usage(item_id: int, u: EquipmentUsageIn, db: Session = Depends(get_db)):
    if u.end_time <= u.start_time:
        raise HTTPException(400, "Время окончания должно быть позже начала")
    item = db.execute(text("SELECT id, name, item_type FROM items WHERE id = :id"), {"id": item_id}).fetchone()
    if not item:
        raise HTTPException(404, "Прибор не найден")
    if item.item_type != "equipment":
        raise HTTPException(400, "Позиция не является оборудованием")

    # Проверка ТО
    on_maintenance = db.execute(text("""
SELECT id, reason FROM maintenance
WHERE item_id = :item_id AND date_from <= :d AND date_to >= :d
LIMIT 1
    """), {"item_id": item_id, "d": u.start_time.date()}).fetchone()
    if on_maintenance:
        reason = f": {on_maintenance.reason}" if on_maintenance.reason else ""
        raise HTTPException(409, f"Прибор на техобслуживании{reason}")

    # Проверка конфликта времени
    conflict = db.execute(text("""
SELECT id, user_name, notes FROM equipment_usage
WHERE item_id = :item_id
  AND tsrange(start_time, end_time) && tsrange(:start_time, :end_time)
LIMIT 1
    """), {"item_id": item_id, "start_time": u.start_time, "end_time": u.end_time}).fetchone()

    if conflict and not u.force:
        raise HTTPException(409, f"CONFLICT:{conflict.user_name}:{conflict.notes or ''}")

    try:
        row = db.execute(
            text("""
INSERT INTO equipment_usage (item_id, user_name, start_time, end_time, notes)
VALUES (:item_id, :user_name, :start_time, :end_time, :notes)
RETURNING id, item_id, user_name, start_time, end_time, notes, created_at
            """),
            {"item_id": item_id, "user_name": u.user_name, "start_time": u.start_time, "end_time": u.end_time, "notes": u.notes},
        ).fetchone()
    except IntegrityError:
        db.rollback()
        raise HTTPException(409, "Конфликт бронирования")
    log_activity(db, "забронировано", "equipment_usage", row[0], item[1] or str(item_id), f"{u.user_name} {u.start_time}–{u.end_time}")
    db.commit()
    return dict(row._mapping)

@app.delete("/api/equipment/usage/{usage_id}")
def delete_equipment_usage(usage_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT item_id, user_name FROM equipment_usage WHERE id = :id"), {"id": usage_id}).fetchone()
    cur = db.execute(text("DELETE FROM equipment_usage WHERE id = :id"), {"id": usage_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    if row:
        log_activity(db, "отменено", "equipment_usage", usage_id, row[1], f"item_id={row[0]}")
    db.commit()
    return {"ok": True}

# ── Maintenance (ТО) ─────────────────────────────────────────────────────────

@app.get("/api/equipment/{item_id}/maintenance")
def list_maintenance(item_id: int, db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT id, item_id, date_from, date_to, reason, created_by, created_at
FROM maintenance
WHERE item_id = :item_id
ORDER BY date_from
    """), {"item_id": item_id}).fetchall()
    return [dict(r._mapping) for r in rows]

@app.get("/api/maintenance/all")
def list_all_maintenance(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT m.id, m.item_id, m.date_from, m.date_to, m.reason, m.created_by, m.created_at,
       i.name as item_name
FROM maintenance m
JOIN items i ON m.item_id = i.id
ORDER BY m.date_from
    """)).fetchall()
    return [dict(r._mapping) for r in rows]

@app.post("/api/equipment/{item_id}/maintenance")
def add_maintenance(item_id: int, m: MaintenanceIn, db: Session = Depends(get_db)):
    item = db.execute(text("SELECT id, name, item_type FROM items WHERE id = :id"), {"id": item_id}).fetchone()
    if not item:
        raise HTTPException(404, "Прибор не найден")
    if item.item_type != "equipment":
        raise HTTPException(400, "Позиция не является оборудованием")
    if m.date_to < m.date_from:
        raise HTTPException(400, "Дата окончания раньше начала")
    row = db.execute(
        text("""
INSERT INTO maintenance (item_id, date_from, date_to, reason, created_by)
VALUES (:item_id, :date_from, :date_to, :reason, :created_by)
RETURNING id, item_id, date_from, date_to, reason, created_by, created_at
        """),
        {"item_id": item_id, "date_from": m.date_from, "date_to": m.date_to, "reason": m.reason, "created_by": m.created_by},
    ).fetchone()
    log_activity(db, "ТО добавлено", "maintenance", row[0], item.name or f"item_id={item_id}", f"{m.date_from}–{m.date_to}" + (f" {m.reason}" if m.reason else ""))
    db.commit()
    return dict(row._mapping)

@app.delete("/api/equipment/maintenance/{maintenance_id}")
def delete_maintenance(maintenance_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT item_id, reason FROM maintenance WHERE id = :id"), {"id": maintenance_id}).fetchone()
    cur = db.execute(text("DELETE FROM maintenance WHERE id = :id"), {"id": maintenance_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    if row:
        log_activity(db, "ТО удалено", "maintenance", maintenance_id, f"item_id={row[0]}", row[1] or "")
    db.commit()
    return {"ok": True}

# ── Клеточная культура (боксы) ────────────────────────────────────────────────

@app.get("/api/cell-usage")
def list_cell_usage(
    period: str = Query(default="week"),
    offset: int = Query(default=0),
    db: Session = Depends(get_db),
):
    start, end = period_range(period, offset)
    rows = db.execute(text("""
SELECT id, box_number, user_name, start_time, end_time, experiment, created_at
FROM cell_usage
WHERE start_time >= :start AND end_time <= :end
ORDER BY start_time
    """), {"start": start, "end": end}).fetchall()
    return {
        "period": period,
        "offset": offset,
        "range_start": start,
        "range_end": end,
        "results": [dict(r._mapping) for r in rows],
    }

@app.post("/api/cell-usage")
def add_cell_usage(u: CellUsageIn, db: Session = Depends(get_db)):
    if u.end_time <= u.start_time:
        raise HTTPException(400, "Время окончания должно быть позже начала")
    if u.box_number not in (1, 2):
        raise HTTPException(400, "Допустимые боксы: 1 или 2")

    # Проверка конфликта
    conflict = db.execute(text("""
SELECT id, user_name, experiment FROM cell_usage
WHERE box_number = :box_number
  AND tsrange(start_time, end_time) && tsrange(:start_time, :end_time)
LIMIT 1
    """), {"box_number": u.box_number, "start_time": u.start_time, "end_time": u.end_time}).fetchone()

    if conflict and not u.force:
        raise HTTPException(409, f"CONFLICT:{conflict.user_name}:{conflict.experiment or ''}")

    row = db.execute(
        text("""
INSERT INTO cell_usage (box_number, user_name, start_time, end_time, experiment)
VALUES (:box_number, :user_name, :start_time, :end_time, :experiment)
RETURNING id, box_number, user_name, start_time, end_time, experiment, created_at
        """),
        {"box_number": u.box_number, "user_name": u.user_name, "start_time": u.start_time, "end_time": u.end_time, "experiment": u.experiment},
    ).fetchone()
    db.commit()
    return dict(row._mapping)

@app.delete("/api/cell-usage/{usage_id}")
def delete_cell_usage(usage_id: int, db: Session = Depends(get_db)):
    cur = db.execute(text("DELETE FROM cell_usage WHERE id = :id"), {"id": usage_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    db.commit()
    return {"ok": True}




@app.patch("/api/cell-usage/{usage_id}")
def update_cell_usage(usage_id: int, data: CellUsageUpdate, db: Session = Depends(get_db)):
    updates = {k: v for k, v in data.dict().items() if v is not None}
    if not updates:
        raise HTTPException(400, "Нет данных")
    sets = ", ".join(f"{k} = :{k}" for k in updates)
    updates["id"] = usage_id
    cur = db.execute(text(f"UPDATE cell_usage SET {sets} WHERE id = :id"), updates)
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    db.commit()
    return {"ok": True}





# ── Обработка экспериментов: Alamar Blue ─────────────────────────────────────


# ── Эндпоинты: расписание приборов (instrument_usage) ──────────────────────
@app.get("/api/instrument-usage")
def list_instrument_usage(instrument_id: str = Query(default=""), period: str = Query(default="week"), offset: int = Query(default=0), db: Session = Depends(get_db)):
    start, end = period_range(period, offset)
    params = {"start": start, "end": end}
    where = "start_time < :end AND end_time > :start"
    if instrument_id:
        where += " AND instrument_id = :instrument_id"
        params["instrument_id"] = instrument_id
    rows = db.execute(text(f"""
        SELECT id, instrument_id, username, start_time, end_time, notes, created_at
        FROM instrument_usage
        WHERE {where}
        ORDER BY start_time
    """), params).fetchall()
    return {"period": period, "offset": offset, "range_start": start, "range_end": end,
            "results": [dict(r._mapping) for r in rows]}


@app.post("/api/instrument-usage")
def add_instrument_usage(u: InstrumentUsageIn, db: Session = Depends(get_db)):
    if u.end_time <= u.start_time:
        raise HTTPException(400, "Время окончания должно быть позже начала")

    on_maintenance = db.execute(text("""
        SELECT id, reason FROM instrument_maintenance
        WHERE instrument_id = :instrument_id AND date_from <= :d AND date_to >= :d
        LIMIT 1
    """), {"instrument_id": u.instrument_id, "d": u.start_time.date()}).fetchone()
    if on_maintenance:
        reason = f" ({on_maintenance.reason})" if on_maintenance.reason else ""
        raise HTTPException(409, f"Прибор на техобслуживании{reason}")

    conflict = db.execute(text("""
        SELECT id, username, notes FROM instrument_usage
        WHERE instrument_id = :instrument_id
          AND tsrange(start_time, end_time) && tsrange(:start_time, :end_time)
        LIMIT 1
    """), {"instrument_id": u.instrument_id, "start_time": u.start_time, "end_time": u.end_time}).fetchone()
    if conflict and not u.force:
        raise HTTPException(409, f"КОНФЛИКТ: {conflict.username} ({conflict.notes or ''})")

    row = db.execute(text("""
        INSERT INTO instrument_usage (instrument_id, username, start_time, end_time, notes)
        VALUES (:instrument_id, :username, :start_time, :end_time, :notes)
        RETURNING id, instrument_id, username, start_time, end_time, notes, created_at
    """), {"instrument_id": u.instrument_id, "username": u.username, "start_time": u.start_time,
           "end_time": u.end_time, "notes": u.notes}).fetchone()

    log_activity(db, "add", "instrument_usage", row[0], u.instrument_id, f"{u.username} {u.start_time}-{u.end_time}")
    db.commit()
    return dict(row._mapping)


@app.patch("/api/instrument-usage/{usage_id}")
def update_instrument_usage(usage_id: int, data: InstrumentUsageUpdateIn, db: Session = Depends(get_db)):
    updates = {k: v for k, v in data.dict().items() if v is not None}
    if not updates:
        raise HTTPException(400, "Нет данных для обновления")
    sets = ", ".join(f"{k} = :{k}" for k in updates)
    updates["id"] = usage_id
    cur = db.execute(text(f"UPDATE instrument_usage SET {sets} WHERE id = :id"), updates)
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    db.commit()
    return {"ok": True}


@app.delete("/api/instrument-usage/{usage_id}")
def delete_instrument_usage(usage_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT instrument_id, username FROM instrument_usage WHERE id = :id"), {"id": usage_id}).fetchone()
    cur = db.execute(text("DELETE FROM instrument_usage WHERE id = :id"), {"id": usage_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    if row:
        log_activity(db, "delete", "instrument_usage", usage_id, row.username, f"instrument_id={row.instrument_id}")
    db.commit()
    return {"ok": True}


@app.get("/api/instrument-maintenance/all")
def list_all_instrument_maintenance(db: Session = Depends(get_db)):
    rows = db.execute(text("""
        SELECT id, instrument_id, date_from, date_to, reason, created_by, created_at
        FROM instrument_maintenance
        ORDER BY date_from
    """)).fetchall()
    return [dict(r._mapping) for r in rows]


@app.post("/api/instrument-maintenance")
def add_instrument_maintenance(m: InstrumentMaintenanceIn, db: Session = Depends(get_db)):
    if m.date_to < m.date_from:
        raise HTTPException(400, "Дата окончания должна быть не раньше даты начала")
    row = db.execute(text("""
        INSERT INTO instrument_maintenance (instrument_id, date_from, date_to, reason, created_by)
        VALUES (:instrument_id, :date_from, :date_to, :reason, :created_by)
        RETURNING id, instrument_id, date_from, date_to, reason, created_by, created_at
    """), {"instrument_id": m.instrument_id, "date_from": m.date_from, "date_to": m.date_to,
           "reason": m.reason, "created_by": m.created_by}).fetchone()
    log_activity(db, "add", "instrument_maintenance", row[0], m.instrument_id,
                 f"{m.date_from}-{m.date_to} {m.reason or ''}")
    db.commit()
    return dict(row._mapping)


@app.delete("/api/instrument-maintenance/{maintenance_id}")
def delete_instrument_maintenance(maintenance_id: int, db: Session = Depends(get_db)):
    row = db.execute(text("SELECT instrument_id, reason FROM instrument_maintenance WHERE id = :id"), {"id": maintenance_id}).fetchone()
    cur = db.execute(text("DELETE FROM instrument_maintenance WHERE id = :id"), {"id": maintenance_id})
    if cur.rowcount == 0:
        raise HTTPException(404, "Не найдено")
    if row:
        log_activity(db, "delete", "instrument_maintenance", maintenance_id, row.instrument_id, row.reason or "")
    db.commit()
    return {"ok": True}
# ── END Эндпоинты: расписание приборов ──────────────────────────────────────


@app.post("/api/experiments/alamar-blue")
async def process_alamar_blue(file: UploadFile = File(...)):
    """
    Обработка Alamar Blue по логике файла пользователя и по мотивам Bio-Rad.

    Формулы из пользовательского шаблона:
    corrected = OD570 * K570 - OD600 * K600

    Для расчёта жизнеспособности используются два варианта:
    1) relative_to_positive_control = corrected / positive_control_corrected * 100
    2) relative_to_negative_control = corrected / negative_control_corrected * 100

    Где:
    - K570 и K600 берутся из первой строки Excel-шаблона
    - negative control = среднее по B:I последней строки 600 нм
    - positive control = среднее по J:M последней строки 600 нм
    """
    try:
        import openpyxl
        content = await file.read()
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
        ws = wb[wb.sheetnames[0]]

        def f(cell):
            v = ws[cell].value
            return None if v is None else float(v)

        k570 = float(ws['J1'].value or 117216)
        k600 = float(ws['H1'].value or 80586)

        # 570nm block: rows 4..11, cols B..M
        od570 = []
        for r in range(4, 12):
            row = []
            for c in range(2, 14):
                val = ws.cell(r, c).value
                row.append(float(val) if val is not None else None)
            od570.append(row)

        # 600nm block: rows 15..22, cols B..M
        od600 = []
        for r in range(15, 23):
            row = []
            for c in range(2, 14):
                val = ws.cell(r, c).value
                row.append(float(val) if val is not None else None)
            od600.append(row)

        corrected = []
        for r in range(8):
            crow = []
            for c in range(12):
                a = od570[r][c]
                b = od600[r][c]
                if a is None or b is None:
                    crow.append(None)
                else:
                    crow.append(round(a * k570 - b * k600, 6))
            corrected.append(crow)

        neg_ctrl_600 = [od600[7][c] for c in range(0, 8) if od600[7][c] is not None]
        pos_ctrl_600 = [od600[7][c] for c in range(8, 12) if od600[7][c] is not None]
        neg_ctrl_570 = [od570[7][c] for c in range(0, 8) if od570[7][c] is not None]
        pos_ctrl_570 = [od570[7][c] for c in range(8, 12) if od570[7][c] is not None]

        avg_neg_600 = sum(neg_ctrl_600) / len(neg_ctrl_600) if neg_ctrl_600 else 0
        avg_pos_600 = sum(pos_ctrl_600) / len(pos_ctrl_600) if pos_ctrl_600 else 0
        avg_neg_570 = sum(neg_ctrl_570) / len(neg_ctrl_570) if neg_ctrl_570 else 0
        avg_pos_570 = sum(pos_ctrl_570) / len(pos_ctrl_570) if pos_ctrl_570 else 0

        neg_corrected = avg_neg_570 * k570 - avg_neg_600 * k600
        pos_corrected = avg_pos_570 * k570 - avg_pos_600 * k600

        viability_neg = []
        viability_pos = []
        for r in range(8):
            row_neg = []
            row_pos = []
            for c in range(12):
                v = corrected[r][c]
                row_neg.append(round(v / neg_corrected * 100, 2) if v is not None and neg_corrected else None)
                row_pos.append(round(v / pos_corrected * 100, 2) if v is not None and pos_corrected else None)
            viability_neg.append(row_neg)
            viability_pos.append(row_pos)

        # Группировка: 4 образца по 3 колонки
        sample_names = ['Образец 1', 'Образец 2', 'Образец 3', 'Образец 4']
        samples = []
        for i, start_col in enumerate([0, 3, 6, 9]):
            vals_pos = []
            vals_neg = []
            for r in range(8):
                for c in range(start_col, min(start_col + 3, 12)):
                    if viability_pos[r][c] is not None:
                        vals_pos.append(viability_pos[r][c])
                    if viability_neg[r][c] is not None:
                        vals_neg.append(viability_neg[r][c])
            if vals_pos:
                mean_pos = round(sum(vals_pos) / len(vals_pos), 2)
                std_pos = round((sum((x - mean_pos) ** 2 for x in vals_pos) / len(vals_pos)) ** 0.5, 2)
                mean_neg = round(sum(vals_neg) / len(vals_neg), 2) if vals_neg else None
                std_neg = round((sum((x - mean_neg) ** 2 for x in vals_neg) / len(vals_neg)) ** 0.5, 2) if vals_neg else None
                samples.append({
                    'name': sample_names[i] if i < len(sample_names) else f'Образец {i+1}',
                    'mean_positive_control': mean_pos,
                    'std_positive_control': std_pos,
                    'mean_negative_control': mean_neg,
                    'std_negative_control': std_neg,
                    'n': len(vals_pos),
                })

        return {
            'ok': True,
            'sheet_name': ws.title,
            'k570': k570,
            'k600': k600,
            'od570': od570,
            'od600': od600,
            'corrected': corrected,
            'negative_control_corrected': round(neg_corrected, 6),
            'positive_control_corrected': round(pos_corrected, 6),
            'avg_negative_570': round(avg_neg_570, 6),
            'avg_negative_600': round(avg_neg_600, 6),
            'avg_positive_570': round(avg_pos_570, 6),
            'avg_positive_600': round(avg_pos_600, 6),
            'viability_negative_control': viability_neg,
            'viability_positive_control': viability_pos,
            'row_labels': list('ABCDEFGH'),
            'samples': samples,
            'formula_corrected': 'corrected = OD570 * K570 - OD600 * K600',
            'formula_negative': '% viability (negative control) = corrected / negative_control_corrected * 100',
            'formula_positive': '% viability (positive control) = corrected / positive_control_corrected * 100',
            'reference': 'Bio-Rad alamarBlue spectrophotometric approach + user Excel template logic',
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f'Ошибка обработки файла: {str(e)}')

# ── Обработка экспериментов: Release / Загрузка ──────────────────────────────

@app.post("/api/experiments/release")
async def process_release(
    file: UploadFile = File(...),
    sheet_name: Optional[str] = Form(default=None),
):
    """
    Обработка загрузки и релиза по логике пользовательского Excel-шаблона.
    Поддерживает несколько листов; если sheet_name не указан, берётся первый.
    Все коэффициенты читаются из Excel, вручную вводить их не нужно.
    """
    try:
        import openpyxl
        content = await file.read()
        wb = openpyxl.load_workbook(BytesIO(content), data_only=True)
        ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb[wb.sheetnames[0]]

        def fv(row, col):
            val = ws.cell(row=row, column=col).value
            return None if val is None else float(val)

        # Блок загрузки
        solvent_loading = ws['A2'].value
        k_loading = fv(2, 2)
        b_loading = fv(2, 3)
        pack_mass_mg = fv(4, 2)
        loading_abs = fv(7, 2)
        loading_dilution = fv(7, 3)
        loading_vol_initial = fv(7, 4)
        loading_vol_result = fv(7, 5)

        conc_loading = ((loading_abs - b_loading) / k_loading) if (loading_abs is not None and k_loading) else None
        mass_solution = conc_loading * loading_vol_result if conc_loading is not None and loading_vol_result is not None else None
        percent_solution = mass_solution / pack_mass_mg * 100 if mass_solution is not None and pack_mass_mg else None
        percent_loaded = 100 - percent_solution if percent_solution is not None else None
        encapsulated_mass = pack_mass_mg * percent_loaded / 100 if percent_loaded is not None and pack_mass_mg is not None else None
        release_pack_mass = encapsulated_mass / 2 if encapsulated_mass is not None else None

        # Блок PBS pH5 / release 1
        solvent_release_1 = ws['A34'].value
        k_release_1 = fv(3, 2)
        b_release_1 = fv(3, 3)
        release_rows_1 = []
        previous_remaining = 100.0
        for r in range(35, 41):
            time_label = ws.cell(r, 1).value
            if time_label is None:
                continue
            reps = [fv(r, 2), fv(r, 3), fv(r, 4)]
            reps = [x for x in reps if x is not None]
            od_mean = sum(reps) / len(reps) if reps else None
            dilution = fv(r, 6)
            vol_initial = fv(r, 7)
            vol_result = fv(r, 8)
            conc = ((od_mean - b_release_1) / k_release_1) if (od_mean is not None and k_release_1) else None
            mass = conc * vol_result if conc is not None and vol_result is not None else None
            pct_in_solution = mass / release_pack_mass * 100 if mass is not None and release_pack_mass else None
            pct_remaining = previous_remaining - pct_in_solution if pct_in_solution is not None else None
            pct_release = 100 - pct_remaining if pct_remaining is not None else None
            release_fraction = pct_release * 0.1 if pct_release is not None else None
            if pct_remaining is not None:
                previous_remaining = pct_remaining
            release_rows_1.append({
                'time': str(time_label),
                'replicates': reps,
                'od_mean': round(od_mean, 6) if od_mean is not None else None,
                'dilution': dilution,
                'volume_initial_ml': vol_initial,
                'volume_result_ml': vol_result,
                'concentration_mg_ml': round(conc, 9) if conc is not None else None,
                'mass_mg': round(mass, 9) if mass is not None else None,
                'percent_in_solution': round(pct_in_solution, 9) if pct_in_solution is not None else None,
                'percent_remaining': round(pct_remaining, 9) if pct_remaining is not None else None,
                'percent_release': round(pct_release, 9) if pct_release is not None else None,
                'release_fraction': round(release_fraction, 9) if release_fraction is not None else None,
            })

        # Блок NaCl pH7 / release 2
        solvent_release_2 = ws['A61'].value
        k_release_2 = fv(43, 2)
        b_release_2 = fv(43, 3)
        release_rows_2 = []
        previous_remaining2 = 100.0
        for r in range(62, 68):
            time_label = ws.cell(r, 1).value
            if time_label is None:
                continue
            reps = [fv(r, 2), fv(r, 3), fv(r, 4)]
            reps = [x for x in reps if x is not None]
            od_mean = sum(reps) / len(reps) if reps else None
            dilution = fv(r, 6)
            vol_initial = fv(r, 7)
            vol_result = fv(r, 8)
            conc = ((od_mean - b_release_2) / k_release_2) if (od_mean is not None and k_release_2) else None
            mass = conc * vol_result if conc is not None and vol_result is not None else None
            pct_in_solution = mass / release_pack_mass * 100 if mass is not None and release_pack_mass else None
            pct_remaining = previous_remaining2 - pct_in_solution if pct_in_solution is not None else None
            pct_release = 100 - pct_remaining if pct_remaining is not None else None
            release_fraction = pct_release * 0.1 if pct_release is not None else None
            if pct_remaining is not None:
                previous_remaining2 = pct_remaining
            release_rows_2.append({
                'time': str(time_label),
                'replicates': reps,
                'od_mean': round(od_mean, 6) if od_mean is not None else None,
                'dilution': dilution,
                'volume_initial_ml': vol_initial,
                'volume_result_ml': vol_result,
                'concentration_mg_ml': round(conc, 9) if conc is not None else None,
                'mass_mg': round(mass, 9) if mass is not None else None,
                'percent_in_solution': round(pct_in_solution, 9) if pct_in_solution is not None else None,
                'percent_remaining': round(pct_remaining, 9) if pct_remaining is not None else None,
                'percent_release': round(pct_release, 9) if pct_release is not None else None,
                'release_fraction': round(release_fraction, 9) if release_fraction is not None else None,
            })

        return {
            'ok': True,
            'sheet_name': ws.title,
            'available_sheets': wb.sheetnames,
            'loading': {
                'solvent': solvent_loading,
                'k': k_loading,
                'b': b_loading,
                'pack_mass_mg': pack_mass_mg,
                'absorbance_423': loading_abs,
                'dilution': loading_dilution,
                'volume_initial_ml': loading_vol_initial,
                'volume_result_ml': loading_vol_result,
                'concentration_mg_ml': round(conc_loading, 9) if conc_loading is not None else None,
                'mass_in_solution_mg': round(mass_solution, 9) if mass_solution is not None else None,
                'percent_in_solution': round(percent_solution, 9) if percent_solution is not None else None,
                'percent_loaded': round(percent_loaded, 9) if percent_loaded is not None else None,
                'encapsulated_mass_mg': round(encapsulated_mass, 9) if encapsulated_mass is not None else None,
                'release_pack_mass_mg': round(release_pack_mass, 9) if release_pack_mass is not None else None,
            },
            'release_profiles': [
                {
                    'solvent': solvent_release_1,
                    'k': k_release_1,
                    'b': b_release_1,
                    'rows': release_rows_1,
                },
                {
                    'solvent': solvent_release_2,
                    'k': k_release_2,
                    'b': b_release_2,
                    'rows': release_rows_2,
                }
            ],
            'formula_loading': 'C = (A - b) / k; mass = C * V; %loaded = 100 - (mass_solution / pack_mass) * 100',
            'formula_release': '%in_solution = mass / release_pack_mass * 100; %remaining(n) = %remaining(n-1) - %in_solution; %release = 100 - %remaining',
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f'Ошибка обработки файла: {str(e)}')

@app.post("/api/experiments/release/export-excel")
async def export_release_excel(body: dict):
    """Экспортирует результаты обработки релиза в Excel."""
    try:
        wb = Workbook()

        # Лист загрузки
        ws1 = wb.active
        ws1.title = "Загрузка"
        loading = body.get("loading", {})
        if loading and "error" not in loading:
            ws1.append(["Параметр", "Значение"])
            ws1.append(["Общая масса (мкг)", loading.get("mass_total_ug", "")])
            ws1.append(["Масса в супернатанте (мкг)", loading.get("mass_in_supernatant_ug", "")])
            ws1.append(["Масса запакованная (мкг)", loading.get("mass_encapsulated_ug", "")])
            ws1.append(["Эффективность загрузки (%)", loading.get("encapsulation_efficiency_pct", "")])
            ws1.append(["Концентрация супернатанта (мкг/мл)", loading.get("concentration_supernatant_ug_ml", "")])
            ws1.append(["", ""])
            ws1.append(["Absorbance (измерения)"] + loading.get("absorbance_values", []))

        # Лист релиза
        profiles = body.get("release_profiles", [])
        for profile in profiles:
            if not profile.get('rows'):
                continue
            name = str(profile.get('solvent') or 'Релиз')[:31]
            ws2 = wb.create_sheet(name)
            ws2.append(["Время", "OD mean", "Конц. мг/мл", "Масса мг", "% в растворе", "% осталось", "% релиза", "доля релиза"])
            for row in profile.get('rows', []):
                ws2.append([row.get('time'), row.get('od_mean'), row.get('concentration_mg_ml'), row.get('mass_mg'), row.get('percent_in_solution'), row.get('percent_remaining'), row.get('percent_release'), row.get('release_fraction')])

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=release_results.xlsx"},
        )
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/experiments/alamar-blue/export-excel")
async def export_alamar_excel(body: dict):
    """Экспортирует результаты Alamar Blue в Excel."""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Жизнеспособность"

        samples = body.get("samples", [])
        if samples:
            ws.append(["Образец", "Mean vs +control (%)", "Std vs +control (%)", "Mean vs -control (%)", "Std vs -control (%)", "n"])
            for s in samples:
                ws.append([s.get("name"), s.get("mean_positive_control"), s.get("std_positive_control"), s.get("mean_negative_control"), s.get("std_negative_control"), s.get("n")])

        ws2 = wb.create_sheet("Планшет (% viability)")
        row_labels = body.get("row_labels", list("ABCDEFGH"))
        ws2.append([""] + [str(c + 1) for c in range(12)])
        viab = body.get("viability_positive_control", [])
        for r_idx, row_data in enumerate(viab):
            label = row_labels[r_idx] if r_idx < len(row_labels) else str(r_idx + 1)
            ws2.append([label] + [v if v is not None else "" for v in row_data])

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=alamar_blue_results.xlsx"},
        )
    except Exception as e:
        raise HTTPException(500, str(e))

# ── Export ────────────────────────────────────────────────────────────────────

@app.get("/api/export/reagents/excel")
def export_reagents_excel(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT i.internal_code, i.name, i.formula, i.cas, i.manufacturer,
       i.catalog_number, i.inventory_number, i.quantity, i.unit,
       l.room, l.cabinet
FROM items i LEFT JOIN locations l ON i.location_id = l.id
WHERE i.item_type = 'reagent' ORDER BY l.room, l.cabinet, i.name
    """)).fetchall()
    headers = ["Код", "Название", "Формула", "CAS", "Производитель", "Каталожный №", "Инв. №", "Количество", "Ед.", "Комната", "Шкаф"]
    return _export_excel(rows, headers, "reagents.xlsx")

@app.get("/api/export/reagents/word")
def export_reagents_word(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT i.internal_code, i.name, i.formula, i.cas, l.room, l.cabinet
FROM items i LEFT JOIN locations l ON i.location_id = l.id
WHERE i.item_type = 'reagent' ORDER BY l.room, l.cabinet, i.name
    """)).fetchall()
    headers = ["Код", "Название", "Формула", "CAS", "Комната", "Шкаф"]
    return _export_word_table(rows, headers, "Реактивы", "reagents.docx")

@app.get("/api/export/equipment/excel")
def export_equipment_excel(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT i.internal_code, i.name, i.manufacturer, i.serial_number,
       i.inventory_number, l.room, l.cabinet, i.document_url, i.gdrive_link
FROM items i LEFT JOIN locations l ON i.location_id = l.id
WHERE i.item_type = 'equipment' ORDER BY l.room, l.cabinet, i.name
    """)).fetchall()
    headers = ["Код", "Название", "Производитель", "Зав. №", "Инв. №", "Комната", "Шкаф", "Документ", "Google Drive"]
    return _export_excel(rows, headers, "equipment.xlsx")

@app.get("/api/export/equipment/word")
def export_equipment_word(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT i.internal_code, i.name, i.manufacturer, l.room, l.cabinet
FROM items i LEFT JOIN locations l ON i.location_id = l.id
WHERE i.item_type = 'equipment' ORDER BY l.room, l.cabinet, i.name
    """)).fetchall()
    headers = ["Код", "Название", "Производитель", "Комната", "Шкаф"]
    return _export_word_table(rows, headers, "Оборудование", "equipment.docx")

@app.get("/api/export/consumables/excel")
def export_consumables_excel(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT i.internal_code, i.name, i.manufacturer, i.catalog_number,
       i.quantity, i.unit, l.room, l.cabinet
FROM items i LEFT JOIN locations l ON i.location_id = l.id
WHERE i.item_type = 'consumable' ORDER BY l.room, l.cabinet, i.name
    """)).fetchall()
    headers = ["Код", "Название", "Производитель", "Каталожный №", "Количество", "Ед.", "Комната", "Шкаф"]
    return _export_excel(rows, headers, "consumables.xlsx")

@app.get("/api/export/consumables/word")
def export_consumables_word(db: Session = Depends(get_db)):
    rows = db.execute(text("""
SELECT i.internal_code, i.name, i.quantity, i.unit, l.room, l.cabinet
FROM items i LEFT JOIN locations l ON i.location_id = l.id
WHERE i.item_type = 'consumable' ORDER BY l.room, l.cabinet, i.name
    """)).fetchall()
    headers = ["Код", "Название", "Количество", "Ед.", "Комната", "Шкаф"]
    return _export_word_table(rows, headers, "Расходники", "consumables.docx")

@app.get("/api/export/protocols/excel")
def export_protocols_excel(db: Session = Depends(get_db)):
    rows = db.execute(text("SELECT code, name, category, steps FROM protocols ORDER BY category, name")).fetchall()
    headers = ["Код", "Название", "Категория", "Этапы синтеза"]
    return _export_excel(rows, headers, "protocols.xlsx")

@app.get("/api/export/protocols/word")
def export_protocols_word(db: Session = Depends(get_db)):
    rows = db.execute(text("SELECT code, name, category, steps FROM protocols ORDER BY category, name")).fetchall()
    doc = Document()
    doc.add_heading("Протоколы синтезов", level=1)
    for code, name, category, steps in rows:
        doc.add_heading(f"{code} — {name}" if code else name, level=2)
        doc.add_paragraph(f"Категория: {category}")
        for line in (steps or "").split("\n"):
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=protocols.docx"},
    )

@app.get("/api/export/sql-template")
def export_sql_template():
    return {"message": "Use import_xlsx.py to load Excel files into items."}

# ── Static files (SPA) ────────────────────────────────────────────────────────

app.mount("/", StaticFiles(directory="static", html=True), name="static")
